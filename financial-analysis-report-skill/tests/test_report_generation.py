from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from render_report import build_report  # noqa: E402
from validate_report import validate_docx  # noqa: E402


@pytest.fixture()
def generated_report(tmp_path: Path) -> Path:
    data = json.loads((ROOT / "examples" / "sample_financial_data.json").read_text(encoding="utf-8"))
    output = tmp_path / "report.docx"
    build_report(data, output)
    return output


def _border_value(cell, edge: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        return None
    node = tc_borders.find(qn(f"w:{edge}"))
    if node is None:
        return None
    return node.get(qn("w:val"))


def test_body_style_uses_required_fonts_and_size(generated_report: Path):
    doc = Document(generated_report)
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size == Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    assert rfonts.get(qn("w:eastAsia")) == "SimSun"
    assert rfonts.get(qn("w:ascii")) == "Times New Roman"
    assert rfonts.get(qn("w:hAnsi")) == "Times New Roman"


def test_body_paragraphs_have_first_line_indent_and_one_point_five_spacing(generated_report: Path):
    doc = Document(generated_report)
    body = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    assert body, "report must contain normal body paragraphs"
    for paragraph in body[:8]:
        assert paragraph.paragraph_format.first_line_indent == Pt(24)
        assert paragraph.paragraph_format.line_spacing == pytest.approx(1.5)


def test_report_contains_required_sections_and_multiple_tables(generated_report: Path):
    doc = Document(generated_report)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for heading in [
        "一、核心结论",
        "二、总体经营与财务表现",
        "三、盈利能力分析",
        "四、偿债能力与资本结构",
        "五、营运效率分析",
        "六、现金流量分析",
        "七、主要风险与关注事项",
        "八、管理建议",
        "附录：指标口径与数据来源",
    ]:
        assert heading in full_text
    assert len(doc.tables) >= 5


def test_tables_use_black_three_line_structure_without_vertical_borders(generated_report: Path):
    doc = Document(generated_report)
    for table in doc.tables:
        assert len(table.rows) >= 2
        first_cell = table.rows[0].cells[0]
        last_cell = table.rows[-1].cells[0]
        assert _border_value(first_cell, "top") == "single"
        assert _border_value(first_cell, "bottom") == "single"
        assert _border_value(last_cell, "bottom") == "single"
        for row in table.rows:
            for cell in row.cells:
                assert _border_value(cell, "left") in (None, "nil")
                assert _border_value(cell, "right") in (None, "nil")


def test_structural_validator_passes_generated_report(generated_report: Path):
    result = validate_docx(generated_report)
    assert result["status"] == "pass", result
    assert result["summary"]["table_count"] >= 5
    assert result["summary"]["body_paragraph_count"] >= 8
