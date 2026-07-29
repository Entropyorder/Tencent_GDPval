from __future__ import annotations

from dataclasses import dataclass
from decimal import Decimal, ROUND_HALF_UP
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


@dataclass(frozen=True)
class ReportStyle:
    page_top_cm: float = 2.6
    page_bottom_cm: float = 2.4
    page_left_cm: float = 3.0
    page_right_cm: float = 2.6
    body_cn_font: str = "SimSun"
    body_en_font: str = "Times New Roman"
    body_size_pt: float = 12.0
    table_size_pt: float = 10.5
    first_line_indent_pt: float = 24.0
    title_size_pt: float = 22.0
    h1_size_pt: float = 16.0
    h2_size_pt: float = 14.0
    line_spacing: float = 1.5
    table_top_sz: str = "12"     # eighths of a point = 1.5 pt
    table_mid_sz: str = "6"      # 0.75 pt
    table_bottom_sz: str = "12"  # 1.5 pt


STYLE = ReportStyle()


def set_run_fonts(run, *, cn: str, en: str, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = en
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:cs"), en)


def set_style_fonts(style, *, cn: str, en: str, size_pt: float, bold: bool | None = None) -> None:
    style.font.name = en
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:cs"), en)


def configure_document(document: Document, *, company_short_name: str, report_title: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(STYLE.page_top_cm)
    section.bottom_margin = Cm(STYLE.page_bottom_cm)
    section.left_margin = Cm(STYLE.page_left_cm)
    section.right_margin = Cm(STYLE.page_right_cm)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    set_style_fonts(normal, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.body_size_pt)
    normal.paragraph_format.first_line_indent = Pt(STYLE.first_line_indent_pt)
    normal.paragraph_format.line_spacing = STYLE.line_spacing
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _upsert_paragraph_style(document, "Report Title", STYLE.title_size_pt, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=0, space_after=18)
    _upsert_paragraph_style(document, "Report Subtitle", 14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=0, space_after=8)
    _upsert_paragraph_style(document, "Heading 1", STYLE.h1_size_pt, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, space_before=14, space_after=8, keep_with_next=True)
    _upsert_paragraph_style(document, "Heading 2", STYLE.h2_size_pt, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, space_before=10, space_after=6, keep_with_next=True)
    _upsert_paragraph_style(document, "Table Title", STYLE.table_size_pt, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=8, space_after=4, keep_with_next=True)
    _upsert_paragraph_style(document, "Source Note", 9, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=0, space_before=3, space_after=8)
    _upsert_paragraph_style(document, "Cover Meta", 12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0, space_before=0, space_after=8)
    _upsert_paragraph_style(document, "Small Note", 9, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=18, space_before=0, space_after=0)

    _configure_header_footer(section, company_short_name=company_short_name, report_title=report_title)
    _configure_doc_defaults(document)


def _upsert_paragraph_style(
    document: Document,
    name: str,
    size_pt: float,
    *,
    bold: bool,
    alignment,
    first_indent: float,
    space_before: float,
    space_after: float,
    keep_with_next: bool = False,
) -> None:
    styles = document.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_style_fonts(style, cn="SimHei" if name in {"Report Title", "Heading 1", "Heading 2"} else STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=size_pt, bold=bold)
    fmt = style.paragraph_format
    fmt.alignment = alignment
    fmt.first_line_indent = Pt(first_indent)
    fmt.line_spacing = STYLE.line_spacing if name not in {"Report Title", "Report Subtitle", "Table Title", "Cover Meta", "Source Note"} else 1.0
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.keep_with_next = keep_with_next


def _configure_doc_defaults(document: Document) -> None:
    styles_el = document.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), STYLE.body_en_font)
    rfonts.set(qn("w:hAnsi"), STYLE.body_en_font)
    rfonts.set(qn("w:eastAsia"), STYLE.body_cn_font)
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), str(int(STYLE.body_size_pt * 2)))


def _configure_header_footer(section, *, company_short_name: str, report_title: str) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{company_short_name}｜{report_title}")
    set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=9)
    _set_paragraph_bottom_border(p, size="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Pt(0)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("第 ")
    set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=9)
    _append_field(run, "PAGE")
    run2 = fp.add_run(" 页")
    set_run_fonts(run2, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=9)


def _append_field(run, field_code: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _set_paragraph_bottom_border(paragraph, *, size: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)


def add_body_paragraph(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Pt(STYLE.first_line_indent_pt)
    p.paragraph_format.line_spacing = STYLE.line_spacing
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_fonts(lead, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.body_size_pt, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_fonts(rest, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.body_size_pt)
    else:
        run = p.add_run(text)
        set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.body_size_pt)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = document.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_fonts(run, cn="SimHei", en=STYLE.body_en_font, size_pt=STYLE.h1_size_pt if level == 1 else STYLE.h2_size_pt, bold=True)


def add_table_title(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Table Title")
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.table_size_pt, bold=True)


def add_source_note(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Source Note")
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=9)


def add_three_line_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    numeric_columns: Iterable[int] = (),
    widths_cm: Sequence[float] | None = None,
) -> None:
    numeric_columns = set(numeric_columns)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = Pt(22)
    _set_repeat_table_header(table.rows[0])

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        if widths_cm:
            cell.width = Cm(widths_cm[col_idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(header))
        set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.table_size_pt, bold=True)

    for row_values in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(21)
        for col_idx, value in enumerate(row_values):
            cell = row.cells[col_idx]
            if widths_cm:
                cell.width = Cm(widths_cm[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_fonts(run, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=STYLE.table_size_pt)

    _apply_three_line_borders(table)
    _set_table_cell_margins(table, top=70, start=90, bottom=70, end=90)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _apply_three_line_borders(table) -> None:
    row_count = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                old = tc_borders.find(qn(f"w:{edge}"))
                if old is not None:
                    tc_borders.remove(old)

            top_val = "single" if r_idx == 0 else "nil"
            bottom_val = "single" if r_idx in {0, row_count - 1} else "nil"
            _append_cell_border(tc_borders, "top", top_val, STYLE.table_top_sz if r_idx == 0 else "0")
            _append_cell_border(tc_borders, "bottom", bottom_val, STYLE.table_mid_sz if r_idx == 0 else (STYLE.table_bottom_sz if r_idx == row_count - 1 else "0"))
            _append_cell_border(tc_borders, "left", "nil", "0")
            _append_cell_border(tc_borders, "right", "nil", "0")
            _append_cell_border(tc_borders, "insideH", "nil", "0")
            _append_cell_border(tc_borders, "insideV", "nil", "0")
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None:
                shd.set(qn("w:fill"), "FFFFFF")


def _append_cell_border(tc_borders, edge: str, val: str, size: str) -> None:
    node = OxmlElement(f"w:{edge}")
    node.set(qn("w:val"), val)
    node.set(qn("w:sz"), size)
    node.set(qn("w:space"), "0")
    node.set(qn("w:color"), "000000")
    tc_borders.append(node)


def _set_table_cell_margins(table, *, top: int, start: int, bottom: int, end: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def format_amount(value: float | int, decimals: int = 0) -> str:
    q = Decimal(str(value)).quantize(Decimal("1") if decimals == 0 else Decimal("1." + "0" * decimals), rounding=ROUND_HALF_UP)
    if q < 0:
        return f"({abs(q):,.{decimals}f})"
    return f"{q:,.{decimals}f}"


def format_percent(value: float | None, decimals: int = 1) -> str:
    if value is None:
        return "不可比"
    q = Decimal(str(value * 100)).quantize(Decimal("1." + "0" * decimals), rounding=ROUND_HALF_UP)
    return f"{q:.{decimals}f}%"


def format_ratio(value: float | None, decimals: int = 2) -> str:
    if value is None:
        return "不可比"
    q = Decimal(str(value)).quantize(Decimal("1." + "0" * decimals), rounding=ROUND_HALF_UP)
    return f"{q:.{decimals}f}"
