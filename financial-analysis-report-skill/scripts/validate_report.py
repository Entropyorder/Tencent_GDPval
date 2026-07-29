from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from report_style import STYLE

REQUIRED_HEADINGS = [
    "一、核心结论",
    "二、总体经营与财务表现",
    "三、盈利能力分析",
    "四、偿债能力与资本结构",
    "五、营运效率分析",
    "六、现金流量分析",
    "七、主要风险与关注事项",
    "八、管理建议",
    "附录：指标口径与数据来源",
]


def _border_val(cell, edge: str) -> str | None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        return None
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        return None
    return node.get(qn("w:val"))


def validate_docx(path: str | Path) -> dict[str, Any]:
    doc = Document(path)
    errors: list[str] = []
    warnings: list[str] = []

    normal = doc.styles["Normal"]
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    east_asia = rfonts.get(qn("w:eastAsia")) if rfonts is not None else None
    ascii_font = rfonts.get(qn("w:ascii")) if rfonts is not None else None
    hansi_font = rfonts.get(qn("w:hAnsi")) if rfonts is not None else None
    if east_asia != STYLE.body_cn_font:
        errors.append(f"Normal 中文字体应为 {STYLE.body_cn_font}，实际为 {east_asia}")
    if ascii_font != STYLE.body_en_font or hansi_font != STYLE.body_en_font:
        errors.append(f"Normal 英文字体应为 {STYLE.body_en_font}")
    if normal.font.size != Pt(STYLE.body_size_pt):
        errors.append(f"Normal 字号应为 {STYLE.body_size_pt} pt")

    body_paragraphs = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    for index, paragraph in enumerate(body_paragraphs, start=1):
        if paragraph.paragraph_format.first_line_indent != Pt(STYLE.first_line_indent_pt):
            errors.append(f"正文第 {index} 段首行缩进不为 {STYLE.first_line_indent_pt} pt")
        spacing = paragraph.paragraph_format.line_spacing
        if spacing is None or abs(float(spacing) - STYLE.line_spacing) > 0.01:
            errors.append(f"正文第 {index} 段行距不为 {STYLE.line_spacing} 倍")

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for heading in REQUIRED_HEADINGS:
        if heading not in full_text:
            errors.append(f"缺少章节：{heading}")

    if len(doc.tables) < 5:
        errors.append("三线表数量少于 5 个")

    for table_index, table in enumerate(doc.tables, start=1):
        if len(table.rows) < 2:
            errors.append(f"表 {table_index} 行数不足")
            continue
        for cell in table.rows[0].cells:
            if _border_val(cell, "top") != "single":
                errors.append(f"表 {table_index} 缺少顶线")
            if _border_val(cell, "bottom") != "single":
                errors.append(f"表 {table_index} 缺少表头下线")
        for cell in table.rows[-1].cells:
            if _border_val(cell, "bottom") != "single":
                errors.append(f"表 {table_index} 缺少底线")
        for row in table.rows:
            for cell in row.cells:
                if _border_val(cell, "left") not in (None, "nil") or _border_val(cell, "right") not in (None, "nil"):
                    errors.append(f"表 {table_index} 存在竖线")
                    break
        if any(len(cell.text.strip()) == 0 for cell in table.rows[0].cells):
            warnings.append(f"表 {table_index} 存在空表头")

    status = "pass" if not errors else "fail"
    return {
        "status": status,
        "summary": {
            "paragraph_count": len(doc.paragraphs),
            "body_paragraph_count": len(body_paragraphs),
            "table_count": len(doc.tables),
            "error_count": len(errors),
            "warning_count": len(warnings),
        },
        "errors": errors,
        "warnings": warnings,
        "checks": {
            "body_font_cn": east_asia,
            "body_font_en": ascii_font,
            "body_size_pt": normal.font.size.pt if normal.font.size else None,
            "first_line_indent_pt": STYLE.first_line_indent_pt,
            "line_spacing": STYLE.line_spacing,
            "required_headings": REQUIRED_HEADINGS,
            "three_line_tables": len(doc.tables),
        },
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="校验财务分析报告 DOCX 的结构与格式")
    parser.add_argument("docx", help="待校验 DOCX")
    parser.add_argument("--output", help="校验结果 JSON 输出路径")
    args = parser.parse_args()
    result = validate_docx(args.docx)
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(payload, encoding="utf-8")
    print(payload)
    raise SystemExit(0 if result["status"] == "pass" else 1)


if __name__ == "__main__":
    main()
