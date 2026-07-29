from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from report_style import (
    STYLE,
    add_body_paragraph,
    add_heading,
    add_source_note,
    add_table_title,
    add_three_line_table,
    configure_document,
    format_amount,
    format_percent,
    format_ratio,
    set_run_fonts,
)

REQUIRED_TOP_LEVEL = {
    "company",
    "report",
    "unit",
    "current_period",
    "prior_period",
    "management_context",
}

REQUIRED_FINANCIAL_FIELDS = {
    "revenue",
    "cost",
    "gross_profit",
    "selling_expense",
    "administrative_expense",
    "rd_expense",
    "finance_expense",
    "operating_profit",
    "net_profit",
    "total_assets",
    "current_assets",
    "cash",
    "accounts_receivable",
    "inventory",
    "total_liabilities",
    "current_liabilities",
    "interest_bearing_debt",
    "equity",
    "operating_cash_flow",
    "investing_cash_flow",
    "financing_cash_flow",
    "capital_expenditure",
}


def _require(data: Mapping[str, Any], path: str) -> Any:
    current: Any = data
    for part in path.split("."):
        if not isinstance(current, Mapping) or part not in current:
            raise ValueError(f"缺少必填字段：{path}")
        current = current[part]
    return current


def validate_input(data: Mapping[str, Any]) -> None:
    missing = sorted(REQUIRED_TOP_LEVEL - set(data))
    if missing:
        raise ValueError("缺少顶层字段：" + ", ".join(missing))
    for period_name in ("current_period", "prior_period"):
        period = _require(data, period_name)
        absent = sorted(REQUIRED_FINANCIAL_FIELDS - set(period))
        if absent:
            raise ValueError(f"{period_name} 缺少字段：" + ", ".join(absent))
    if _require(data, "current_period.label") == _require(data, "prior_period.label"):
        raise ValueError("本期与上期标签不得相同")


def safe_div(numerator: float, denominator: float) -> float | None:
    return None if denominator == 0 else numerator / denominator


def calculate_metrics(current: Mapping[str, float], prior: Mapping[str, float]) -> dict[str, dict[str, float | None]]:
    avg_assets = (current["total_assets"] + prior["total_assets"]) / 2
    avg_ar = (current["accounts_receivable"] + prior["accounts_receivable"]) / 2
    avg_inventory = (current["inventory"] + prior["inventory"]) / 2
    return {
        "current": {
            "revenue_growth": safe_div(current["revenue"] - prior["revenue"], prior["revenue"]),
            "gross_margin": safe_div(current["gross_profit"], current["revenue"]),
            "operating_margin": safe_div(current["operating_profit"], current["revenue"]),
            "net_margin": safe_div(current["net_profit"], current["revenue"]),
            "roa": safe_div(current["net_profit"], avg_assets),
            "roe": safe_div(current["net_profit"], (current["equity"] + prior["equity"]) / 2),
            "current_ratio": safe_div(current["current_assets"], current["current_liabilities"]),
            "quick_ratio": safe_div(current["current_assets"] - current["inventory"], current["current_liabilities"]),
            "debt_ratio": safe_div(current["total_liabilities"], current["total_assets"]),
            "debt_to_equity": safe_div(current["interest_bearing_debt"], current["equity"]),
            "ar_turnover": safe_div(current["revenue"], avg_ar),
            "inventory_turnover": safe_div(current["cost"], avg_inventory),
            "cfo_net_profit": safe_div(current["operating_cash_flow"], current["net_profit"]),
            "free_cash_flow": current["operating_cash_flow"] - current["capital_expenditure"],
        },
        "prior": {
            "revenue_growth": None,
            "gross_margin": safe_div(prior["gross_profit"], prior["revenue"]),
            "operating_margin": safe_div(prior["operating_profit"], prior["revenue"]),
            "net_margin": safe_div(prior["net_profit"], prior["revenue"]),
            "roa": safe_div(prior["net_profit"], prior["total_assets"]),
            "roe": safe_div(prior["net_profit"], prior["equity"]),
            "current_ratio": safe_div(prior["current_assets"], prior["current_liabilities"]),
            "quick_ratio": safe_div(prior["current_assets"] - prior["inventory"], prior["current_liabilities"]),
            "debt_ratio": safe_div(prior["total_liabilities"], prior["total_assets"]),
            "debt_to_equity": safe_div(prior["interest_bearing_debt"], prior["equity"]),
            "ar_turnover": safe_div(prior["revenue"], prior["accounts_receivable"]),
            "inventory_turnover": safe_div(prior["cost"], prior["inventory"]),
            "cfo_net_profit": safe_div(prior["operating_cash_flow"], prior["net_profit"]),
            "free_cash_flow": prior["operating_cash_flow"] - prior["capital_expenditure"],
        },
    }


def _add_cover(document: Document, data: Mapping[str, Any]) -> None:
    company = data["company"]
    report = data["report"]
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph(style="Report Title")
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(report["title"])
    set_run_fonts(run, cn="SimHei", en=STYLE.body_en_font, size_pt=STYLE.title_size_pt, bold=True)

    subtitle = document.add_paragraph(style="Report Subtitle")
    subtitle.paragraph_format.first_line_indent = Pt(0)
    r = subtitle.add_run(f"{report['period']}｜内部管理报告")
    set_run_fonts(r, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=14)

    for _ in range(8):
        document.add_paragraph()
    for line in [
        f"编制单位：{company['name']}",
        f"编制部门：{report['prepared_by']}",
        f"报告日期：{report['date']}",
        f"数据属性：{report.get('data_nature', '内部管理数据')}",
    ]:
        meta = document.add_paragraph(style="Cover Meta")
        meta.paragraph_format.first_line_indent = Pt(0)
        rr = meta.add_run(line)
        set_run_fonts(rr, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=12)

    note = document.add_paragraph(style="Small Note")
    note.paragraph_format.first_line_indent = Pt(0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = note.add_run("本报告仅用于内部经营分析与管理决策，不构成审计意见或对外披露文件。")
    set_run_fonts(rr, cn=STYLE.body_cn_font, en=STYLE.body_en_font, size_pt=9)
    document.add_page_break()


def _metric_change(current: float, prior: float) -> str:
    if prior == 0:
        return "不可比"
    return format_percent((current - prior) / abs(prior))


def _pp_change(current: float | None, prior: float | None) -> str:
    if current is None or prior is None:
        return "不可比"
    return f"{(current - prior) * 100:+.1f}个百分点"


def build_report(data: Mapping[str, Any], output_path: str | Path) -> Path:
    validate_input(data)
    current = data["current_period"]
    prior = data["prior_period"]
    metrics = calculate_metrics(current, prior)
    curm, prm = metrics["current"], metrics["prior"]
    unit = data["unit"]
    company = data["company"]
    report = data["report"]
    context = data["management_context"]

    document = Document()
    configure_document(document, company_short_name=company["short_name"], report_title=report["title"])
    props = document.core_properties
    props.title = report["title"]
    props.subject = "财务分析"
    props.author = report["prepared_by"]
    props.keywords = "财务分析, 盈利能力, 偿债能力, 现金流, 营运效率"
    props.comments = "由 financial-analysis-report skill 生成并经结构化校验。"

    _add_cover(document, data)

    add_heading(document, "一、核心结论", 1)
    add_body_paragraph(
        document,
        f"报告期内，公司营业收入为{format_amount(current['revenue'])}{unit}，较上期增长{format_percent(curm['revenue_growth'])}；净利润为{format_amount(current['net_profit'])}{unit}，同比增长{_metric_change(current['net_profit'], prior['net_profit'])}。收入与利润同步改善，说明本期增长并非单纯依赖规模扩张，产品结构调整和制造效率提升已经开始转化为经营收益。",
    )
    add_body_paragraph(
        document,
        f"盈利质量总体向好。综合毛利率由{format_percent(prm['gross_margin'])}提升至{format_percent(curm['gross_margin'])}，净利率提高{_pp_change(curm['net_margin'], prm['net_margin'])}；经营活动现金净流量达到{format_amount(current['operating_cash_flow'])}{unit}，为净利润的{format_percent(curm['cfo_net_profit'])}。利润与现金流的匹配度较高，但资本性支出仍使自由现金流为{format_amount(curm['free_cash_flow'])}{unit}，扩产项目的投入节奏需要继续与订单兑现情况联动。",
    )
    add_body_paragraph(
        document,
        f"资产负债结构保持稳健，资产负债率为{format_percent(curm['debt_ratio'])}，较上期下降{abs((curm['debt_ratio'] - prm['debt_ratio']) * 100):.1f}个百分点；流动比率和速动比率分别为{format_ratio(curm['current_ratio'])}和{format_ratio(curm['quick_ratio'])}。短期偿债能力处于可控区间，但应收账款余额增长快于部分回款节点，仍需把客户信用分层和合同收款条款作为下一阶段的重点管理事项。",
    )

    add_table_title(document, f"表1  核心财务指标概览（单位：{unit}，比率除外）")
    add_three_line_table(
        document,
        ["指标", current["label"], prior["label"], "变动"],
        [
            ["营业收入", format_amount(current["revenue"]), format_amount(prior["revenue"]), _metric_change(current["revenue"], prior["revenue"])],
            ["营业利润", format_amount(current["operating_profit"]), format_amount(prior["operating_profit"]), _metric_change(current["operating_profit"], prior["operating_profit"])],
            ["净利润", format_amount(current["net_profit"]), format_amount(prior["net_profit"]), _metric_change(current["net_profit"], prior["net_profit"])],
            ["经营活动现金净流量", format_amount(current["operating_cash_flow"]), format_amount(prior["operating_cash_flow"]), _metric_change(current["operating_cash_flow"], prior["operating_cash_flow"])],
            ["资产负债率", format_percent(curm["debt_ratio"]), format_percent(prm["debt_ratio"]), _pp_change(curm["debt_ratio"], prm["debt_ratio"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[6.3, 3.0, 3.0, 3.2],
    )
    add_source_note(document, "资料来源：公司财务账套及管理报表；本示例使用模拟数据。")

    add_heading(document, "二、总体经营与财务表现", 1)
    add_body_paragraph(
        document,
        f"本期收入增长主要来自{context['growth_driver']}。其中，核心产品订单保持稳定，新客户导入使产能利用率较上期改善；与此同时，公司主动收缩低毛利、回款周期较长的订单，收入增速因此低于订单增速，但毛利贡献更为扎实。该变化说明经营策略已由“追求规模”逐步转向“兼顾规模、盈利和现金回收”。",
    )
    add_body_paragraph(
        document,
        f"期间费用合计为{format_amount(current['selling_expense'] + current['administrative_expense'] + current['rd_expense'] + current['finance_expense'])}{unit}，同比增长{_metric_change(current['selling_expense'] + current['administrative_expense'] + current['rd_expense'] + current['finance_expense'], prior['selling_expense'] + prior['administrative_expense'] + prior['rd_expense'] + prior['finance_expense'])}，低于营业收入增幅。销售费用基本稳定，管理费用随组织扩张温和增加，研发投入继续向关键工艺和产品平台倾斜。费用率下降并非由削减必要投入形成，而是收入增长对固定费用产生了摊薄作用。",
    )
    add_body_paragraph(
        document,
        f"期末总资产为{format_amount(current['total_assets'])}{unit}，较上期增加{format_amount(current['total_assets'] - prior['total_assets'])}{unit}。新增资产主要投向{context['investment_focus']}，与公司当前产能提升和产品升级方向一致。需要关注的是，资产扩张对未来订单和现金流提出了更高要求，若新增产能爬坡慢于计划，折旧及资金占用将对后续利润率形成压力。",
    )

    add_table_title(document, f"表2  收入、成本与费用结构（单位：{unit}）")
    add_three_line_table(
        document,
        ["项目", current["label"], prior["label"], "同比变动"],
        [
            ["营业收入", format_amount(current["revenue"]), format_amount(prior["revenue"]), _metric_change(current["revenue"], prior["revenue"])],
            ["营业成本", format_amount(current["cost"]), format_amount(prior["cost"]), _metric_change(current["cost"], prior["cost"])],
            ["销售费用", format_amount(current["selling_expense"]), format_amount(prior["selling_expense"]), _metric_change(current["selling_expense"], prior["selling_expense"])],
            ["管理费用", format_amount(current["administrative_expense"]), format_amount(prior["administrative_expense"]), _metric_change(current["administrative_expense"], prior["administrative_expense"])],
            ["研发费用", format_amount(current["rd_expense"]), format_amount(prior["rd_expense"]), _metric_change(current["rd_expense"], prior["rd_expense"])],
            ["财务费用", format_amount(current["finance_expense"]), format_amount(prior["finance_expense"]), _metric_change(current["finance_expense"], prior["finance_expense"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[5.5, 3.2, 3.2, 3.5],
    )
    add_source_note(document, "注：费用项目按公司管理口径列示，未进行非经常性项目调整。")

    add_heading(document, "三、盈利能力分析", 1)
    add_body_paragraph(
        document,
        f"综合毛利率为{format_percent(curm['gross_margin'])}，较上期提升{(curm['gross_margin'] - prm['gross_margin']) * 100:.1f}个百分点。改善主要来自产品组合优化、单位制造成本下降以及采购价格趋稳。虽然人工和折旧成本仍有上升，但产量增长和良率改善对单位成本形成了有效对冲，毛利率提升具有一定经营基础。",
    )
    add_body_paragraph(
        document,
        f"营业利润率和净利率分别为{format_percent(curm['operating_margin'])}和{format_percent(curm['net_margin'])}。净利润增幅高于收入增幅，反映规模效应开始释放。不过，当前利润率仍受融资成本和新项目爬坡费用影响，尚不宜将本期改善简单外推为长期稳定水平。后续判断盈利趋势，应重点观察新产品毛利、设备利用率和客户价格年降三项因素。",
    )
    add_body_paragraph(
        document,
        f"净资产收益率为{format_percent(curm['roe'])}，较上期有所提升。该变化主要由净利率改善驱动，而非明显提高财务杠杆，因此收益质量相对稳健。公司仍应优先通过提高资产周转效率和项目回报率改善股东回报，避免以增加有息负债换取短期指标上升。",
    )

    add_table_title(document, "表3  盈利能力指标")
    add_three_line_table(
        document,
        ["指标", current["label"], prior["label"], "变化"],
        [
            ["综合毛利率", format_percent(curm["gross_margin"]), format_percent(prm["gross_margin"]), _pp_change(curm["gross_margin"], prm["gross_margin"])],
            ["营业利润率", format_percent(curm["operating_margin"]), format_percent(prm["operating_margin"]), _pp_change(curm["operating_margin"], prm["operating_margin"])],
            ["销售净利率", format_percent(curm["net_margin"]), format_percent(prm["net_margin"]), _pp_change(curm["net_margin"], prm["net_margin"])],
            ["总资产收益率", format_percent(curm["roa"]), format_percent(prm["roa"]), _pp_change(curm["roa"], prm["roa"])],
            ["净资产收益率", format_percent(curm["roe"]), format_percent(prm["roe"]), _pp_change(curm["roe"], prm["roe"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[6.0, 3.0, 3.0, 3.4],
    )
    add_source_note(document, "注：本期总资产收益率和净资产收益率以期初期末平均余额计算；上期按期末余额近似计算。")

    add_heading(document, "四、偿债能力与资本结构", 1)
    add_body_paragraph(
        document,
        f"期末负债总额为{format_amount(current['total_liabilities'])}{unit}，资产负债率为{format_percent(curm['debt_ratio'])}，较上期下降{abs((curm['debt_ratio'] - prm['debt_ratio']) * 100):.1f}个百分点。权益增长快于负债增长，使资本结构略有改善。当前杠杆水平与制造业扩产阶段相匹配，但仍应关注有息负债的期限分布，避免短债集中到期与建设支出高峰重叠。",
    )
    add_body_paragraph(
        document,
        f"流动比率为{format_ratio(curm['current_ratio'])}，速动比率为{format_ratio(curm['quick_ratio'])}，均较上期小幅改善。流动资产能够覆盖流动负债，但存货和应收账款合计占流动资产比重较高，账面流动性与实际可动用现金之间仍存在差异。偿债能力评估应同时结合回款进度、存货可变现性和授信额度，而不能仅依赖静态比率。",
    )
    add_body_paragraph(
        document,
        f"有息负债与净资产之比为{format_percent(curm['debt_to_equity'])}。本期融资净流入主要用于设备投资和补充项目建设资金，资金用途与长期资产形成基本匹配。建议继续提高中长期融资占比，并将新增借款与明确的订单、项目里程碑和现金回收计划绑定，以降低资金闲置或期限错配风险。",
    )

    add_table_title(document, "表4  偿债能力与资本结构指标")
    add_three_line_table(
        document,
        ["指标", current["label"], prior["label"], "变化"],
        [
            ["流动比率", format_ratio(curm["current_ratio"]), format_ratio(prm["current_ratio"]), f"{curm['current_ratio'] - prm['current_ratio']:+.2f}"],
            ["速动比率", format_ratio(curm["quick_ratio"]), format_ratio(prm["quick_ratio"]), f"{curm['quick_ratio'] - prm['quick_ratio']:+.2f}"],
            ["资产负债率", format_percent(curm["debt_ratio"]), format_percent(prm["debt_ratio"]), _pp_change(curm["debt_ratio"], prm["debt_ratio"])],
            ["有息负债/净资产", format_percent(curm["debt_to_equity"]), format_percent(prm["debt_to_equity"]), _pp_change(curm["debt_to_equity"], prm["debt_to_equity"])],
            ["期末现金余额", format_amount(current["cash"]), format_amount(prior["cash"]), _metric_change(current["cash"], prior["cash"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[6.0, 3.0, 3.0, 3.4],
    )
    add_source_note(document, "注：速动资产按流动资产扣除存货计算；有息负债包括短期借款、长期借款及其他计息债务。")

    add_heading(document, "五、营运效率分析", 1)
    add_body_paragraph(
        document,
        f"应收账款周转率为{format_ratio(curm['ar_turnover'])}次，高于上期的{format_ratio(prm['ar_turnover'])}次，整体回款效率有所改善。但期末应收账款仍达到{format_amount(current['accounts_receivable'])}{unit}，且部分新增客户处于信用验证期。建议将逾期天数、客户集中度和回款偏差纳入月度经营例会，避免收入增长掩盖账龄结构恶化。",
    )
    add_body_paragraph(
        document,
        f"存货周转率为{format_ratio(curm['inventory_turnover'])}次，较上期提高。周转改善主要来自订单排产准确度提升和原材料安全库存优化，但在制品仍受部分定制项目交付节奏影响。后续应区分战略备料、正常周转库存和呆滞库存，分别设置责任人和处置时限，避免以总量指标掩盖结构性积压。",
    )
    add_body_paragraph(
        document,
        f"从资产占用看，应收账款与存货合计为{format_amount(current['accounts_receivable'] + current['inventory'])}{unit}，占总资产的{format_percent((current['accounts_receivable'] + current['inventory']) / current['total_assets'])}。该比例尚处于可管理范围，但随着业务规模扩大，营运资金需求会同步增加。公司需要将销售预测、采购计划和资金预算形成闭环，减少业务部门与财务部门对现金需求判断的时间差。",
    )

    add_table_title(document, "表5  营运效率指标")
    add_three_line_table(
        document,
        ["指标", current["label"], prior["label"], "变化"],
        [
            ["应收账款周转率", format_ratio(curm["ar_turnover"]), format_ratio(prm["ar_turnover"]), f"{curm['ar_turnover'] - prm['ar_turnover']:+.2f}"],
            ["存货周转率", format_ratio(curm["inventory_turnover"]), format_ratio(prm["inventory_turnover"]), f"{curm['inventory_turnover'] - prm['inventory_turnover']:+.2f}"],
            ["应收账款余额", format_amount(current["accounts_receivable"]), format_amount(prior["accounts_receivable"]), _metric_change(current["accounts_receivable"], prior["accounts_receivable"])],
            ["存货余额", format_amount(current["inventory"]), format_amount(prior["inventory"]), _metric_change(current["inventory"], prior["inventory"])],
            ["应收账款及存货占总资产", format_percent((current["accounts_receivable"] + current["inventory"]) / current["total_assets"]), format_percent((prior["accounts_receivable"] + prior["inventory"]) / prior["total_assets"]), _pp_change((current["accounts_receivable"] + current["inventory"]) / current["total_assets"], (prior["accounts_receivable"] + prior["inventory"]) / prior["total_assets"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[6.2, 3.0, 3.0, 3.2],
    )
    add_source_note(document, "注：本期周转率使用期初期末平均余额；上期周转率按期末余额近似计算。")

    add_heading(document, "六、现金流量分析", 1)
    add_body_paragraph(
        document,
        f"经营活动现金净流量为{format_amount(current['operating_cash_flow'])}{unit}，同比增加{format_amount(current['operating_cash_flow'] - prior['operating_cash_flow'])}{unit}。经营现金流覆盖净利润，说明本期利润具有较好的现金实现基础。现金改善主要来自销售回款增长、采购付款节奏优化以及部分预付款回收，并非通过大幅延迟正常供应商付款形成。",
    )
    add_body_paragraph(
        document,
        f"投资活动现金净流量为{format_amount(current['investing_cash_flow'])}{unit}，主要用于{context['investment_focus']}。资本性支出为{format_amount(current['capital_expenditure'])}{unit}，高于经营现金净流量，自由现金流为{format_amount(curm['free_cash_flow'])}{unit}。当前现金缺口由融资和期初现金共同覆盖，短期内可以承受，但项目建设若出现延期，资金占用周期将被拉长。",
    )
    add_body_paragraph(
        document,
        f"融资活动现金净流量为{format_amount(current['financing_cash_flow'])}{unit}，期末现金余额增至{format_amount(current['cash'])}{unit}。公司仍具备一定流动性缓冲，但现金余额中包含项目专项资金，实际可自由调度金额低于账面余额。建议滚动编制十三周现金流预测，并对大额采购、设备付款和客户回款设置情景边界。",
    )

    document.add_page_break()
    add_table_title(document, f"表6  现金流量概览（单位：{unit}）")
    add_three_line_table(
        document,
        ["项目", current["label"], prior["label"], "变动额"],
        [
            ["经营活动现金净流量", format_amount(current["operating_cash_flow"]), format_amount(prior["operating_cash_flow"]), format_amount(current["operating_cash_flow"] - prior["operating_cash_flow"])],
            ["投资活动现金净流量", format_amount(current["investing_cash_flow"]), format_amount(prior["investing_cash_flow"]), format_amount(current["investing_cash_flow"] - prior["investing_cash_flow"])],
            ["融资活动现金净流量", format_amount(current["financing_cash_flow"]), format_amount(prior["financing_cash_flow"]), format_amount(current["financing_cash_flow"] - prior["financing_cash_flow"])],
            ["资本性支出", format_amount(current["capital_expenditure"]), format_amount(prior["capital_expenditure"]), format_amount(current["capital_expenditure"] - prior["capital_expenditure"])],
            ["自由现金流", format_amount(curm["free_cash_flow"]), format_amount(prm["free_cash_flow"]), format_amount(curm["free_cash_flow"] - prm["free_cash_flow"])],
            ["经营现金/净利润", format_percent(curm["cfo_net_profit"]), format_percent(prm["cfo_net_profit"]), _pp_change(curm["cfo_net_profit"], prm["cfo_net_profit"])],
        ],
        numeric_columns={1, 2, 3},
        widths_cm=[6.2, 3.0, 3.0, 3.2],
    )
    add_source_note(document, "注：自由现金流按经营活动现金净流量减资本性支出计算。括号表示现金净流出或负数。")

    add_heading(document, "七、主要风险与关注事项", 1)
    risk_items = [
        ("回款风险。", f"应收账款余额为{format_amount(current['accounts_receivable'])}{unit}，新增客户和大客户项目回款节奏存在差异。若收入确认与合同回款不同步，经营现金流可能在季度间明显波动。"),
        ("产能投资风险。", f"本期资本性支出为{format_amount(current['capital_expenditure'])}{unit}，新设备尚处于逐步达产阶段。订单兑现、良率爬坡或客户认证延迟，均可能拉长投资回收期。"),
        ("毛利波动风险。", "产品价格年降、关键材料采购价格和产能利用率仍是影响毛利率的核心变量。当前毛利改善具有经营基础，但不应忽视客户议价和行业竞争带来的压力。"),
        ("流动性风险。", "账面现金包含项目专项资金，且未来十二个月仍有设备尾款和债务到期。若回款低于预算，应及时调整非刚性投资和融资安排。"),
    ]
    for lead, body in risk_items:
        add_body_paragraph(document, lead + body, bold_lead=lead)

    add_heading(document, "八、管理建议", 1)
    recommendation_items = [
        ("第一，", "建立“订单—收入—回款—现金”一体化看板。销售、交付和财务使用同一客户维度追踪合同金额、验收节点、开票进度和到账偏差，将逾期回款责任落实到具体客户与项目。"),
        ("第二，", "对资本性支出实行里程碑拨付。设备采购和项目建设应与订单覆盖、客户认证、产能利用率及实际良率挂钩，对未达到前置条件的非关键支出延后执行。"),
        ("第三，", "推进存货分层管理。对战略备料、正常周转、项目在制和呆滞物料分别设定周转目标、减值规则和处置期限，形成采购、生产、销售共同负责的闭环。"),
        ("第四，", "优化债务期限结构。优先以中长期资金匹配设备和厂房投入，控制短期借款集中到期；新增融资应测算利率、担保条件和项目现金回收周期，不以单一资金成本作为决策依据。"),
        ("第五，", "将利润率改善拆解到产品与客户。月度分析应同时观察收入、单位毛利、质量损失、售后成本和资金占用，避免高收入但低回报项目占用关键产能。"),
    ]
    for lead, body in recommendation_items:
        add_body_paragraph(document, lead + body, bold_lead=lead)

    add_heading(document, "附录：指标口径与数据来源", 1)
    add_heading(document, "（一）主要指标口径", 2)
    definitions = [
        "综合毛利率＝（营业收入－营业成本）÷营业收入。",
        "销售净利率＝净利润÷营业收入。",
        "流动比率＝流动资产÷流动负债；速动比率＝（流动资产－存货）÷流动负债。",
        "资产负债率＝负债总额÷资产总额；有息负债/净资产＝有息负债÷所有者权益。",
        "应收账款周转率＝营业收入÷平均应收账款；存货周转率＝营业成本÷平均存货。",
        "自由现金流＝经营活动现金净流量－资本性支出。",
    ]
    for item in definitions:
        add_body_paragraph(document, item)

    add_heading(document, "（二）数据来源与使用限制", 2)
    add_body_paragraph(document, "本报告数据来源于公司财务账套、资金台账及内部管理报表。示例文件中的企业名称和财务数据均为模拟信息，仅用于展示报告结构、语言和格式，不应作为任何真实企业的经营判断依据。")
    add_body_paragraph(document, "本报告属于管理分析文件，部分指标采用管理口径或近似计算，与经审计财务报表可能存在口径差异。对外披露、融资申报或投资决策使用时，应由财务负责人根据正式报表、审计调整及业务事实复核。")

    _prevent_table_row_splitting(document)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def _prevent_table_row_splitting(document: Document) -> None:
    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None:
                tr_pr.append(OxmlElement("w:cantSplit"))


def main() -> None:
    parser = argparse.ArgumentParser(description="生成规范的中文财务分析报告 Word 文档")
    parser.add_argument("--input", required=True, help="财务数据 JSON 路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 路径")
    args = parser.parse_args()
    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    output = build_report(data, args.output)
    print(output)


if __name__ == "__main__":
    main()
