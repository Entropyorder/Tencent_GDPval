#!/usr/bin/env python3
import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
import zipfile

import fitz
import openpyxl
from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DELIVERABLE_PATTERN = re.compile(
    r"(?m)^\s*\d+[.、]\s+`([^`\n]+\.(?:docx|xlsx|pptx|pdf|csv|md|txt))`"
)
PLACEHOLDER_PATTERN = re.compile(
    r"\bTODO\b|待补充|待填写|示例数据|placeholder",
    re.IGNORECASE,
)
NESTED_FORMULA_EQUALS_PATTERN = re.compile(r"=[A-Z][A-Z0-9_.]*\(")
FINANCIAL_SKILL_NAME = "generating-financial-analysis-reports"


def _is_grey_hex(hexstr):
    """RGB 三通道差 <= 20 视为灰阶（黑/灰/白）；差 > 20 为彩色（违规）。

    None/空/auto/非 6 位十六进制视为合法（不判违规），避免误报。
    """
    if not isinstance(hexstr, str):
        return True
    value = hexstr.lstrip("#")
    if len(value) != 6:
        return True
    try:
        r = int(value[0:2], 16)
        g = int(value[2:4], 16)
        b = int(value[4:6], 16)
    except ValueError:
        return True
    return max(r, g, b) - min(r, g, b) <= 20


def _bw_violations_docx(document):
    """扫描 docx 全部单元格底纹 (w:shd w:fill) 与文字颜色 (w:color w:val)，
    非灰即记违规。"""
    violations = []
    for shd in document.element.iter(qn("w:shd")):
        fill = shd.get(qn("w:fill"))
        if fill and not _is_grey_hex(fill):
            violations.append({"kind": "shading", "color": fill})
    for color in document.element.iter(qn("w:color")):
        val = color.get(qn("w:val"))
        if val and val.lower() != "auto" and not _is_grey_hex(val):
            violations.append({"kind": "font", "color": val})
    return violations


def _bw_violations_xlsx(workbook):
    """扫描 xlsx 单元格底纹与字体颜色，非灰即记违规。"""
    violations = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                fill = cell.fill
                if fill is not None and fill.patternType == "solid":
                    rgb = getattr(fill.fgColor, "rgb", None)
                    if isinstance(rgb, str) and not _is_grey_hex(rgb[-6:]):
                        violations.append(
                            {
                                "cell": f"{worksheet.title}!{cell.coordinate}",
                                "kind": "fill",
                                "color": rgb,
                            }
                        )
                font = cell.font
                if font is not None and font.color is not None:
                    rgb = getattr(font.color, "rgb", None)
                    if isinstance(rgb, str) and not _is_grey_hex(rgb[-6:]):
                        violations.append(
                            {
                                "cell": f"{worksheet.title}!{cell.coordinate}",
                                "kind": "font",
                                "color": rgb,
                            }
                        )
    return violations


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        while chunk := stream.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def expected_deliverables(query):
    section = query.split("## 交付要求", 1)
    if len(section) != 2:
        raise ValueError("query does not contain a 交付要求 section")
    filenames = DELIVERABLE_PATTERN.findall(section[1])
    if not 1 <= len(filenames) <= 5:
        raise ValueError(
            f"expected 1-5 named deliverables in query, found {len(filenames)}"
        )
    if len(filenames) != len(set(filenames)):
        raise ValueError("query contains duplicate deliverable filenames")
    return filenames


def inspect_docx(path):
    document = Document(path)
    paragraph_text = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    table_text = [
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    text = "\n".join(paragraph_text + table_text)
    if len(text) < 1000:
        raise ValueError(f"{path.name}: Word content is too short")
    if PLACEHOLDER_PATTERN.search(text):
        raise ValueError(f"{path.name}: contains placeholder content")
    bw = _bw_violations_docx(document)
    return {
        "paragraphs": len(paragraph_text),
        "tables": len(document.tables),
        "text_chars": len(text),
        "bw_violations": bw,
    }


def inspect_xlsx(path):
    workbook = openpyxl.load_workbook(path, data_only=False, read_only=False)
    nonempty_cells = 0
    formulas = 0
    invalid_formulas = []
    text_parts = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                nonempty_cells += 1
                text_parts.append(str(cell.value))
                if cell.data_type == "f":
                    formulas += 1
                    formula_body = str(cell.value)[1:]
                    if (
                        NESTED_FORMULA_EQUALS_PATTERN.search(formula_body)
                        or any(
                            operator in formula_body
                            for operator in ("+=", "-=", "*=", "/=")
                        )
                    ):
                        invalid_formulas.append(
                            f"{worksheet.title}!{cell.coordinate}"
                        )
                elif str(cell.value).startswith("="):
                    invalid_formulas.append(
                        f"{worksheet.title}!{cell.coordinate}(text-as-formula)"
                    )
    sheet_names = workbook.sheetnames
    bw = _bw_violations_xlsx(workbook)
    workbook.close()
    if nonempty_cells < 50:
        raise ValueError(f"{path.name}: Excel content is too sparse")
    if formulas < 10:
        raise ValueError(f"{path.name}: expected at least 10 formulas")
    if invalid_formulas:
        locations = ", ".join(invalid_formulas[:10])
        raise ValueError(
            f"{path.name}: contains syntactically suspicious formulas at {locations}"
        )
    if PLACEHOLDER_PATTERN.search("\n".join(text_parts)):
        raise ValueError(f"{path.name}: contains placeholder content")
    return {
        "worksheets": sheet_names,
        "nonempty_cells": nonempty_cells,
        "formulas": formulas,
        "suspicious_formulas": 0,
        "bw_violations": bw,
    }


def inspect_file(path):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return inspect_docx(path)
    if suffix == ".xlsx":
        return inspect_xlsx(path)
    if suffix == ".pdf":
        with fitz.open(path) as document:
            if document.page_count < 1:
                raise ValueError(f"{path.name}: PDF has no pages")
            return {"pages": document.page_count}
    if suffix == ".pptx":
        with zipfile.ZipFile(path) as archive:
            slides = [
                name
                for name in archive.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ]
        if not slides:
            raise ValueError(f"{path.name}: PowerPoint has no slides")
        return {"slides": len(slides)}
    text = path.read_text(encoding="utf-8")
    if len(text.strip()) < 200:
        raise ValueError(f"{path.name}: text deliverable is too short")
    if PLACEHOLDER_PATTERN.search(text):
        raise ValueError(f"{path.name}: contains placeholder content")
    return {"text_chars": len(text)}


def validate_financial_resource_usage(task_dir, internal_dir, filenames):
    contract_path = task_dir / "final" / "internal" / "financial_resources.json"
    if not contract_path.is_file():
        return None
    contract = json.loads(contract_path.read_text(encoding="utf-8"))
    usage_path = internal_dir / "resource_usage.json"
    if not usage_path.is_file() or usage_path.stat().st_size == 0:
        raise FileNotFoundError(usage_path)
    usage = json.loads(usage_path.read_text(encoding="utf-8"))

    skill = usage.get("skill", {})
    if skill.get("name") != FINANCIAL_SKILL_NAME:
        raise ValueError("resource usage does not record the financial report skill")
    applied_to = set(skill.get("applied_to", []))
    if applied_to != set(filenames):
        raise ValueError(
            "financial report skill must be recorded for every deliverable"
        )
    principles = skill.get("principles", [])
    if not isinstance(principles, list) or not principles:
        raise ValueError("resource usage must list applied financial skill principles")

    templates_by_id = {
        item["id"]: item for item in contract.get("templates", [])
    }
    template_usage = usage.get("templates", [])
    if not isinstance(template_usage, list) or not template_usage:
        raise ValueError("resource usage must record at least one real template")
    office_deliverables = {
        filename
        for filename in filenames
        if Path(filename).suffix.lower()
        in {".docx", ".xlsx", ".pdf", ".pptx", ".csv"}
    }
    template_coverage = set()
    for item in template_usage:
        resource_id = item.get("id")
        if resource_id not in templates_by_id:
            raise ValueError(
                f"unknown financial template in resource usage: {resource_id}"
            )
        targets = set(item.get("applied_to", []))
        if not targets or not targets.issubset(set(filenames)):
            raise ValueError(
                f"invalid applied_to for financial template: {resource_id}"
            )
        allowed_formats = set(
            templates_by_id[resource_id].get("golden_formats", [])
        )
        for filename in targets:
            file_format = Path(filename).suffix.lower().removeprefix(".")
            if file_format not in allowed_formats:
                raise ValueError(
                    f"template {resource_id} is incompatible with {filename}"
                )
        adaptations = item.get("adaptations", [])
        if not isinstance(adaptations, list) or not adaptations:
            raise ValueError(f"template {resource_id} lacks recorded adaptations")
        if item.get("copied_source_data") is not False:
            raise ValueError(
                f"template {resource_id} must declare copied_source_data=false"
            )
        template_coverage.update(targets)
    if not office_deliverables.issubset(template_coverage):
        missing = sorted(office_deliverables - template_coverage)
        raise ValueError(
            f"deliverables lack a compatible template reference: {missing}"
        )
    return {
        "skill": FINANCIAL_SKILL_NAME,
        "templates": len(template_usage),
        "covered_deliverables": len(template_coverage),
    }


def validate(task_dir, strict_bw="warn"):
    final_dir = task_dir / "final"
    golden_dir = task_dir / "golden solution"
    query = (final_dir / "query.md").read_text(encoding="utf-8")
    filenames = expected_deliverables(query)

    if not golden_dir.is_dir():
        raise FileNotFoundError(golden_dir)
    actual_files = sorted(
        path.name for path in golden_dir.iterdir() if path.is_file()
    )
    if sorted(filenames) != actual_files:
        raise ValueError(
            "golden solution files do not match query deliverables: "
            f"expected={sorted(filenames)} actual={actual_files}"
        )

    inspections = {}
    for filename in filenames:
        path = golden_dir / filename
        inspections[filename] = {
            "bytes": path.stat().st_size,
            "sha256": sha256(path),
            **inspect_file(path),
        }

    internal_dir = golden_dir / "internal"
    for filename in (
        "source_traceability.md",
        "validation_report.json",
        "solution_manifest.json",
    ):
        path = internal_dir / filename
        if not path.is_file() or path.stat().st_size == 0:
            raise FileNotFoundError(path)

    manifest = json.loads(
        (internal_dir / "solution_manifest.json").read_text(encoding="utf-8")
    )
    records = manifest.get("deliverables", [])
    if {record.get("filename") for record in records} != set(filenames):
        raise ValueError("solution manifest deliverables do not match query")
    by_name = {record["filename"]: record for record in records}
    for filename, inspection in inspections.items():
        record = by_name[filename]
        if record.get("sha256") != inspection["sha256"]:
            raise ValueError(f"{filename}: solution manifest SHA-256 mismatch")
        if record.get("bytes") != inspection["bytes"]:
            raise ValueError(f"{filename}: solution manifest byte size mismatch")

    financial_resources = validate_financial_resource_usage(
        task_dir, internal_dir, filenames
    )

    # B/W 后置检查：扫描所有 docx/xlsx 的非灰底纹与字体颜色。
    bw_violations = {}
    for filename, inspection in inspections.items():
        bw = inspection.get("bw_violations", [])
        if bw:
            bw_violations[filename] = bw
    if bw_violations:
        summary = "; ".join(
            f"{name}:{len(items)}" for name, items in bw_violations.items()
        )
        detail = json.dumps(bw_violations, ensure_ascii=False)
        if strict_bw == "error":
            raise ValueError(
                f"non-greyscale (B/W) violations found: {summary}\n{detail}"
            )
        print(f"[warn] non-greyscale (B/W) violations: {summary}\n{detail}", file=sys.stderr)

    return {
        "task": task_dir.name,
        "deliverables": len(filenames),
        "files": inspections,
        "bw_violations": bw_violations,
        "financial_resources": financial_resources,
    }


def main():
    parser = argparse.ArgumentParser(
        description="Validate generated Golden Solution files."
    )
    parser.add_argument("--task", required=True)
    parser.add_argument(
        "--workspace",
        type=Path,
        default=PROJECT_ROOT / "output" / "tasks",
    )
    parser.add_argument(
        "--strict-bw",
        choices=("error", "warn"),
        default="warn",
        help="非灰（彩色）底纹/字体违规的处理：warn 只告警不失败，error 记为失败。",
    )
    args = parser.parse_args()
    if not re.fullmatch(r"\d{3}", args.task):
        raise SystemExit("--task must be a three-digit ID")

    result = validate(args.workspace / f"task_{args.task}", strict_bw=args.strict_bw)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
