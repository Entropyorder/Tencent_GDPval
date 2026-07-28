#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把二进制 Office/PDF 附件解码成文本预览, 落到 golden solution/_extracted/_att_<idx>.txt。

为什么需要它: claude -p 的 Read 工具读 .xlsx/.docx 等二进制很慢且贵,
事实抽取会 600s 超时。先用 Python(openpyxl/python-docx/pypdf)解码成文本,
claude 读文本预览秒级完成, 且更确定(单元格值是 Python 精确读出的)。

既可被 run_golden.py import 调 preextract(case_dir), 也可 CLI 单跑:
    python preextract_attachments.py --case <case_dir> [--max-rows 2000]
"""
import os, sys, json, argparse

def read_csv_text(path):
    """CSV 编码自动探测(UTF-8-sig/UTF-8/GBK/GB18030)。"""
    raw = open(path, "rb").read()
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("gb18030", errors="replace")

def _extract_one(fp, on_disk, max_rows):
    """返回 (kind, lines)。失败在 lines 里记占位, 不抛。"""
    low = on_disk.lower()
    lines = []
    kind = "unknown"
    try:
        if low.endswith(".xlsx") or low.endswith(".xls"):
            kind = "xlsx"
            from openpyxl import load_workbook
            wb = load_workbook(fp, read_only=True, data_only=True)
            for sname in wb.sheetnames:
                ws = wb[sname]
                lines.append(f"[SHEET: {sname}] dims={ws.max_row}x{ws.max_column}")
                rn = 0
                for row in ws.iter_rows(values_only=True):
                    if rn >= max_rows: break
                    cells = ["" if v is None else str(v) for v in row]
                    lines.append(" | ".join(cells))
                    rn += 1
                if rn >= max_rows:
                    lines.append(f"... (truncated at {max_rows} rows)")
            wb.close()
        elif low.endswith(".docx"):
            kind = "docx"
            from docx import Document
            doc = Document(fp)
            pi = 0
            for para in doc.paragraphs:
                if pi >= max_rows: break
                t = para.text.strip()
                if t:
                    lines.append(f"[P] {t}")
                    pi += 1
            for ti, tbl in enumerate(doc.tables):
                lines.append(f"[TABLE {ti+1}] {len(tbl.rows)}x{len(tbl.columns)}")
                rn = 0
                for r in tbl.rows:
                    if rn >= max_rows: break
                    cells = [c.text.strip() for c in r.cells]
                    lines.append(" | ".join(cells))
                    rn += 1
        elif low.endswith(".csv"):
            kind = "csv"
            text = read_csv_text(fp)
            rn = 0
            for ln in text.splitlines():
                if rn >= max_rows: break
                lines.append(ln)
                rn += 1
            if rn >= max_rows:
                lines.append(f"... (truncated at {max_rows} rows)")
        elif low.endswith(".pdf"):
            kind = "pdf"
            try:
                from pypdf import PdfReader
                rd = PdfReader(fp)
                for pno, page in enumerate(rd.pages[:60]):
                    lines.append(f"[PAGE {pno+1}]")
                    lines.append(page.extract_text() or "")
            except Exception as e:
                lines.append(f"(PDF 解析失败: {e})")
        else:
            kind = "text"
            try:
                lines.append(open(fp, encoding="utf-8").read()[:200000])
            except Exception:
                lines.append(read_csv_text(fp)[:200000])
    except Exception as e:
        lines.append(f"(预抽取失败: {e})")
    return kind, lines

def list_attachments(case_dir):
    """扫 case 目录的附件(兼容 GDPval final/attachments 和裸 attachments)。"""
    cands = []
    final_att = os.path.join(case_dir, "final", "attachments")
    bare_att = os.path.join(case_dir, "attachments")
    for base in (final_att, bare_att):
        if os.path.isdir(base):
            cands.append(base)
            break
    if not cands:
        # 退化: case_dir 自身是 attachments 目录
        if os.path.isfile(os.path.join(case_dir, "query.md")) or \
           os.path.isfile(os.path.join(case_dir, "..", "query.md")):
            pass
    atts = []
    for base in cands:
        for f in sorted(os.listdir(base)):
            if f == "__MACOSX" or f.startswith("._") or f.startswith("~$"):
                continue
            fp = os.path.join(base, f)
            if os.path.isfile(fp):
                atts.append({"on_disk": f, "abs": fp, "size": os.path.getsize(fp)})
        break
    return atts

def preextract(case_dir, max_rows=2000):
    """主入口: 解码 case_dir 下附件 → golden solution/_extracted/_att_<i>.txt。
    返回 [{on_disk, real_name, abs, preview_path, size, kind, preview_head}]。"""
    # golden solution/ 目录: GDPval 用 "golden solution", 裸 case 用 golden_solution
    gs = os.path.join(case_dir, "golden solution")
    if not os.path.isdir(gs):
        gs = os.path.join(case_dir, "golden_solution")
    extracted = os.path.join(gs, "_extracted")
    os.makedirs(extracted, exist_ok=True)
    atts = list_attachments(case_dir)
    out = []
    for idx, a in enumerate(atts):
        fp = a["abs"]; on_disk = a["on_disk"]
        kind, lines = _extract_one(fp, on_disk, max_rows)
        header = f"=== FILE idx={idx} name={on_disk} size={a['size']}B kind={kind} ==="
        text = header + "\n" + "\n".join(lines)
        preview_path = os.path.join(extracted, f"_att_{idx}.txt")
        with open(preview_path, "w", encoding="utf-8") as f:
            f.write(text)
        out.append({
            "name": on_disk, "abs": fp,
            "preview_path": preview_path, "size": a["size"], "kind": kind,
            "preview_head": text[:200].replace("\n", " "),
        })
    # 写一个清单方便 agent 索引
    manifest = os.path.join(extracted, "_index.json")
    with open(manifest, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    return out

def main():
    ap = argparse.ArgumentParser(description="预抽取附件 → 文本预览")
    ap.add_argument("--case", required=True, help="case 目录")
    ap.add_argument("--max-rows", type=int, default=2000)
    args = ap.parse_args()
    out = preextract(args.case, args.max_rows)
    print(f"[preextract] {len(out)} 个附件 → golden solution/_extracted/")
    for a in out:
        print(f"  [{a['kind']}] {a['name']}  →  {a['preview_path']}")
    return out

if __name__ == "__main__":
    sys.stdout.reconfigure(line_buffering=True)
    main()
