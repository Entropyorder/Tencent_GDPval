#!/usr/bin/env python3
import argparse
import json
import os
from pathlib import Path
import shutil
import sys

import fitz
import openpyxl
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from finance_forensics.extractors import extract_document


def write_json(path, payload):
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def extract_pdf_full(path):
    chunks = []
    with fitz.open(path) as document:
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").replace("\x00", "").strip()
            if text:
                chunks.append(f"## 第 {page_index} 页\n{text}")
        metadata = {
            "method": "pymupdf-full",
            "total_units": document.page_count,
            "units_read": document.page_count,
        }
    return "\n\n".join(chunks), metadata


def extract_xlsx_full(path):
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    chunks = []
    nonempty_rows = 0
    for sheet_name in workbook.sheetnames:
        lines = [f"## 工作表: {sheet_name}"]
        sheet = workbook[sheet_name]
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = []
            for value in row[:100]:
                text = "" if value is None else str(value).replace("\n", " ").strip()
                values.append(text[:1000])
            while values and not values[-1]:
                values.pop()
            if any(values):
                lines.append(f"{row_index}\t" + "\t".join(values))
                nonempty_rows += 1
            if nonempty_rows >= 10000:
                lines.append("[达到10000个非空行的安全上限]")
                break
        chunks.append("\n".join(lines))
        if nonempty_rows >= 10000:
            break
    total_sheets = len(workbook.sheetnames)
    workbook.close()
    return "\n\n".join(chunks), {
        "method": "openpyxl-full",
        "total_units": total_sheets,
        "units_read": len(chunks),
        "nonempty_rows": nonempty_rows,
    }


def extract_docx_full(path):
    document = Document(path)
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"## 表格 {table_index}")
        for row in table.rows:
            chunks.append(
                "\t".join(cell.text.replace("\n", " ").strip() for cell in row.cells)
            )
    return "\n".join(filter(None, chunks)), {
        "method": "python-docx-full",
        "total_units": len(document.paragraphs) + len(document.tables),
        "units_read": len(document.paragraphs) + len(document.tables),
    }


def extract_for_workspace(path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_full(path)
    if suffix == ".xlsx":
        return extract_xlsx_full(path)
    if suffix == ".docx":
        return extract_docx_full(path)
    result = extract_document(path, max_chars=500_000)
    return result.text, {
        "method": result.method,
        "total_units": result.total_units,
        "units_read": result.units_read,
        "truncated": result.truncated,
        "warnings": result.warnings,
    }


def prepare_task(
    project_root,
    query,
    retrieval_dir=None,
    tasks_dir=None,
    force=False,
):
    query_index = int(query["query_index"])
    task_id = f"{query_index:03d}"
    retrieval_dir = (
        Path(retrieval_dir)
        if retrieval_dir is not None
        else project_root / "output" / "retrieval"
    )
    tasks_dir = (
        Path(tasks_dir)
        if tasks_dir is not None
        else project_root / "output" / "tasks"
    )
    task_dir = tasks_dir / f"task_{task_id}"
    if task_dir.exists():
        if not force:
            raise FileExistsError(f"task workspace already exists: {task_dir}")
        shutil.rmtree(task_dir)

    candidates_dir = task_dir / "candidates"
    extracted_dir = task_dir / "extracted"
    candidates_dir.mkdir(parents=True)
    extracted_dir.mkdir()
    records = []

    for result in query["results"]:
        rank = int(result["rank"])
        filename = result["attachment_filename"]
        source_path = (
            retrieval_dir
            / f"query_{task_id}"
            / "files"
            / filename
        ).resolve()
        if not source_path.is_file():
            raise FileNotFoundError(source_path)

        candidate_name = f"{rank:02d}__{filename}"
        candidate_path = candidates_dir / candidate_name
        os.symlink(source_path, candidate_path)
        text, extraction = extract_for_workspace(source_path)
        extracted_name = f"{rank:02d}__{result['document_id']}.md"
        (extracted_dir / extracted_name).write_text(
            f"# {filename}\n\n"
            f"- rank: {rank}\n"
            f"- document_id: {result['document_id']}\n"
            f"- original_candidate: candidates/{candidate_name}\n"
            f"- extraction: {json.dumps(extraction, ensure_ascii=False)}\n\n"
            f"{text}\n",
            encoding="utf-8",
        )

        record = result.copy()
        record["candidate_path"] = f"candidates/{candidate_name}"
        record["extracted_path"] = f"extracted/{extracted_name}"
        record["workspace_extraction"] = extraction
        records.append(record)

    task_manifest = {
        "schema_version": "1.0",
        "task_id": task_id,
        "candidate_count": len(records),
        "original_query": query["query"],
        "retrieval_query": query["retrieval_query"],
        "candidates": records,
    }
    write_json(task_dir / "candidate_manifest.json", task_manifest)
    (task_dir / "ORIGINAL_QUERY.md").write_text(
        f"# 原始检索查询\n\n{query['query']}\n",
        encoding="utf-8",
    )
    (task_dir / "TASK.md").write_text(
        f"# 题目 {task_id}\n\n"
        f"本工作区只包含查询 {task_id} 对应的 {len(records)} 个候选文件，"
        "不得访问另一道题目的工作区。\n\n"
        "使用十至十七个附件构建一道最终金融题目。必须新增并实际选用一至三个"
        "明确标记的生成附件。严格使用 Pi 的 `gdpval-task-builder` Skill 和"
        "自定义工具，"
        "先固化并重新编号最终附件，再分别反向编写至少十二步的 workflow 和"
        "整体叙述式 query，并完成确定性验收。\n",
        encoding="utf-8",
    )
    return task_dir, len(records)


def main():
    parser = argparse.ArgumentParser(
        description="Prepare isolated Pi workspaces from retrieval results."
    )
    parser.add_argument(
        "--manifest",
        type=Path,
        default=PROJECT_ROOT / "output" / "retrieval" / "manifest.json",
    )
    parser.add_argument(
        "--tasks-dir",
        type=Path,
        default=PROJECT_ROOT / "output" / "tasks",
    )
    parser.add_argument(
        "--task",
        action="append",
        help="Prepare only this three-digit query index; may be repeated.",
    )
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args()

    manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
    queries = manifest.get("queries", [])
    if not queries:
        raise ValueError("retrieval manifest does not contain query candidate sets")

    selected = None
    if args.task:
        if any(not item.isdigit() or len(item) != 3 for item in args.task):
            raise ValueError("--task values must be three-digit query indexes")
        selected = {int(item) for item in args.task}
        available = {int(query["query_index"]) for query in queries}
        missing = selected - available
        if missing:
            raise ValueError(
                f"query indexes not found in manifest: {sorted(missing)}"
            )

    retrieval_dir = args.manifest.resolve().parent
    for query in queries:
        if selected is not None and int(query["query_index"]) not in selected:
            continue
        if len(query.get("results", [])) != 20:
            raise ValueError(
                f"query {query.get('query_index')} must contain exactly 20 candidates"
            )
        task_dir, count = prepare_task(
            PROJECT_ROOT,
            query,
            retrieval_dir=retrieval_dir,
            tasks_dir=args.tasks_dir,
            force=args.force,
        )
        print(f"prepared task={task_dir.name} candidates={count} path={task_dir}")


if __name__ == "__main__":
    main()
