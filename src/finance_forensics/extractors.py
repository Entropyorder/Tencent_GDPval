from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
import csv
import io
import os
import shutil
import subprocess

import fitz
import openpyxl
import xlrd
from charset_normalizer import from_bytes
from docx import Document
from python_calamine import CalamineWorkbook

from .config import PROJECT_ROOT


SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".doc", ".csv"}


@dataclass
class ExtractionResult:
    text: str
    method: str
    total_units: int | None = None
    units_read: int | None = None
    truncated: bool = False
    encoding: str | None = None
    warnings: list[str] = field(default_factory=list)

    @property
    def characters(self):
        return len(self.text)


class ExtractionError(RuntimeError):
    pass


class VisibleHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.chunks = []
        self.ignored_depth = 0

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self.ignored_depth += 1
        elif not self.ignored_depth and tag in {
            "br",
            "div",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "li",
            "p",
            "table",
            "tr",
        }:
            self.chunks.append("\n")
        elif not self.ignored_depth and tag in {"td", "th"}:
            self.chunks.append("\t")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self.ignored_depth = max(0, self.ignored_depth - 1)
        elif not self.ignored_depth and tag in {
            "div",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "li",
            "p",
            "table",
            "tr",
        }:
            self.chunks.append("\n")

    def handle_data(self, data):
        if not self.ignored_depth and data.strip():
            self.chunks.append(data.strip())


def sample_text(text, max_chars):
    text = text.replace("\x00", "").strip()
    if len(text) <= max_chars:
        return text, False
    first = int(max_chars * 0.55)
    middle = int(max_chars * 0.15)
    last = max_chars - first - middle - 80
    center = len(text) // 2
    sampled = (
        text[:first]
        + "\n\n[...中间内容抽样...]\n\n"
        + text[center - middle // 2 : center + middle // 2]
        + "\n\n[...末尾内容抽样...]\n\n"
        + text[-last:]
    )
    return sampled[:max_chars], True


def selected_unit_indexes(total, limit):
    if total <= limit:
        return list(range(total))
    indexes = set(range(min(6, total)))
    slots = max(1, limit - len(indexes))
    for slot in range(1, slots + 1):
        indexes.add(round(slot * (total - 1) / slots))
    return sorted(indexes)[:limit]


def cell_text(value):
    if value is None:
        return ""
    text = str(value).replace("\n", " ").strip()
    return text[:500]


def extract_pdf(path, max_chars):
    chunks = []
    warnings = []
    with fitz.open(path) as document:
        indexes = selected_unit_indexes(document.page_count, 24)
        for page_index in indexes:
            text = document.load_page(page_index).get_text("text")
            if text.strip():
                chunks.append(f"## 第 {page_index + 1} 页\n{text.strip()}")
        total = document.page_count
    text, truncated = sample_text("\n\n".join(chunks), max_chars)
    if len(text) < 200:
        warnings.append("PDF 提取文本很少，可能是扫描件或图片型 PDF")
    return ExtractionResult(
        text=text,
        method="pymupdf",
        total_units=total,
        units_read=len(indexes),
        truncated=truncated,
        warnings=warnings,
    )


def extract_xlsx(path, max_chars):
    try:
        chunks = []
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheet_names = workbook.sheetnames[:12]
        for sheet_name in sheet_names:
            sheet = workbook[sheet_name]
            lines = [f"## 工作表: {sheet_name}"]
            for row_index, row in enumerate(
                sheet.iter_rows(max_row=200, max_col=40, values_only=True), start=1
            ):
                values = [cell_text(value) for value in row]
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    lines.append(f"{row_index}\t" + "\t".join(values))
            chunks.append("\n".join(lines))
        total = len(workbook.sheetnames)
        workbook.close()
        text, truncated = sample_text("\n\n".join(chunks), max_chars)
        return ExtractionResult(
            text=text,
            method="openpyxl",
            total_units=total,
            units_read=len(sheet_names),
            truncated=truncated,
        )
    except Exception as error:
        return extract_spreadsheet_with_calamine(path, max_chars, error)


def extract_xls(path, max_chars):
    try:
        workbook = xlrd.open_workbook(path, on_demand=True)
        chunks = []
        sheet_names = workbook.sheet_names()[:12]
        for sheet_name in sheet_names:
            sheet = workbook.sheet_by_name(sheet_name)
            lines = [f"## 工作表: {sheet_name}"]
            for row_index in range(min(sheet.nrows, 200)):
                values = [
                    cell_text(sheet.cell_value(row_index, column))
                    for column in range(min(sheet.ncols, 40))
                ]
                while values and not values[-1]:
                    values.pop()
                if any(values):
                    lines.append(f"{row_index + 1}\t" + "\t".join(values))
            chunks.append("\n".join(lines))
        total = workbook.nsheets
        workbook.release_resources()
        text, truncated = sample_text("\n\n".join(chunks), max_chars)
        return ExtractionResult(
            text=text,
            method="xlrd",
            total_units=total,
            units_read=len(sheet_names),
            truncated=truncated,
        )
    except Exception as error:
        return extract_spreadsheet_with_calamine(path, max_chars, error)


def extract_spreadsheet_with_calamine(path, max_chars, primary_error):
    workbook = CalamineWorkbook.from_path(path)
    sheet_names = workbook.sheet_names[:12]
    chunks = []
    for sheet_name in sheet_names:
        sheet = workbook.get_sheet_by_name(sheet_name)
        lines = [f"## 工作表: {sheet_name}"]
        if sheet.height == 0:
            chunks.append("\n".join(lines))
            continue
        for row_index, row in enumerate(sheet.iter_rows(), start=1):
            if row_index > 200:
                break
            values = [cell_text(value) for value in row[:40]]
            while values and not values[-1]:
                values.pop()
            if any(values):
                lines.append(f"{row_index}\t" + "\t".join(values))
        chunks.append("\n".join(lines))
    text, truncated = sample_text("\n\n".join(chunks), max_chars)
    return ExtractionResult(
        text=text,
        method="python-calamine",
        total_units=len(workbook.sheet_names),
        units_read=len(sheet_names),
        truncated=truncated,
        warnings=[
            f"主解析器失败，已使用 python-calamine: "
            f"{type(primary_error).__name__}: {primary_error}"
        ],
    )


def extract_docx(path, max_chars):
    document = Document(path)
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs]
    table_count = min(len(document.tables), 30)
    for table_index, table in enumerate(document.tables[:table_count], start=1):
        chunks.append(f"## 表格 {table_index}")
        for row in table.rows[:100]:
            chunks.append("\t".join(cell_text(cell.text) for cell in row.cells))
    text, truncated = sample_text("\n".join(filter(None, chunks)), max_chars)
    return ExtractionResult(
        text=text,
        method="python-docx",
        total_units=len(document.paragraphs) + len(document.tables),
        units_read=len(document.paragraphs) + table_count,
        truncated=truncated,
    )


def antiword_binary():
    candidates = [
        os.environ.get("ANTIWORD_BIN"),
        shutil.which("antiword"),
        str(PROJECT_ROOT / ".tools" / "antiword" / "usr" / "bin" / "antiword"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    return None


def catdoc_binary():
    candidates = [
        os.environ.get("CATDOC_BIN"),
        shutil.which("catdoc"),
        str(PROJECT_ROOT / ".tools" / "catdoc" / "usr" / "bin" / "catdoc"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    return None


def extract_doc_with_catdoc(path, max_chars, antiword_error):
    binary = catdoc_binary()
    if not binary:
        raise ExtractionError(antiword_error)
    environment = os.environ.copy()
    bundled_home = PROJECT_ROOT / ".tools" / "catdoc"
    if (bundled_home / ".catdocrc").is_file():
        environment["HOME"] = str(bundled_home)
    result = subprocess.run(
        [binary, "-d", "utf-8", str(path)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=120,
        check=False,
        env=environment,
    )
    text = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0 or not text.strip():
        error = result.stderr.decode("utf-8", errors="replace").strip()
        raise ExtractionError(
            f"{antiword_error}; catdoc failed with rc={result.returncode}: {error}"
        )
    text, truncated = sample_text(text, max_chars)
    warnings = [f"antiword 失败，已使用 catdoc: {antiword_error}"]
    stderr = result.stderr.decode("utf-8", errors="replace").strip()
    if stderr:
        warnings.append(stderr)
    return ExtractionResult(
        text=text,
        method="catdoc",
        truncated=truncated,
        warnings=warnings,
    )


def extract_doc(path, max_chars):
    binary = antiword_binary()
    if not binary:
        raise ExtractionError(
            "antiword is required for legacy .doc files; configure ANTIWORD_BIN"
        )
    environment = os.environ.copy()
    bundled_resources = PROJECT_ROOT / ".tools" / "antiword" / "usr" / "share" / "antiword"
    if bundled_resources.is_dir():
        environment.setdefault("ANTIWORDHOME", str(bundled_resources))
    result = subprocess.run(
        [binary, "-m", "UTF-8.txt", str(path)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=120,
        check=False,
        env=environment,
    )
    text = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0 and not text.strip():
        error = result.stderr.decode("utf-8", errors="replace").strip()
        antiword_error = f"antiword failed with rc={result.returncode}: {error}"
        return extract_doc_with_catdoc(path, max_chars, antiword_error)
    text, truncated = sample_text(text, max_chars)
    warnings = []
    if result.returncode != 0:
        warnings.append(f"antiword returned rc={result.returncode}")
    return ExtractionResult(
        text=text,
        method="antiword",
        truncated=truncated,
        warnings=warnings,
    )


def extract_csv(path, max_chars):
    byte_limit = 8 * 1024 * 1024
    with path.open("rb") as stream:
        raw = stream.read(byte_limit)
        has_more = bool(stream.read(1))
    match = from_bytes(raw).best()
    if match is None:
        text = raw.decode("gb18030", errors="replace")
        encoding = "gb18030-fallback"
    else:
        text = str(match)
        encoding = match.encoding

    lines = []
    reader = csv.reader(io.StringIO(text))
    for row_index, row in enumerate(reader, start=1):
        if row_index > 400:
            break
        lines.append(f"{row_index}\t" + "\t".join(cell_text(value) for value in row[:50]))
    sampled, truncated = sample_text("\n".join(lines), max_chars)
    return ExtractionResult(
        text=sampled,
        method="csv",
        total_units=None,
        units_read=len(lines),
        truncated=truncated or has_more,
        encoding=encoding,
    )


def extract_html(path, max_chars):
    byte_limit = 8 * 1024 * 1024
    with path.open("rb") as stream:
        raw = stream.read(byte_limit)
        has_more = bool(stream.read(1))
    match = from_bytes(raw).best()
    if match is None:
        source = raw.decode("utf-8", errors="replace")
        encoding = "utf-8-fallback"
    else:
        source = str(match)
        encoding = match.encoding

    parser = VisibleHTMLParser()
    parser.feed(source)
    lines = []
    for line in "".join(parser.chunks).splitlines():
        normalized = " ".join(line.split())
        if normalized:
            lines.append(normalized)
    text, truncated = sample_text("\n".join(lines), max_chars)
    return ExtractionResult(
        text=text,
        method="html-content-sniff",
        truncated=truncated or has_more,
        encoding=encoding,
        warnings=["文件扩展名与实际 HTML 内容不一致"],
    )


def looks_like_html(prefix):
    normalized = prefix.lstrip(b"\xef\xbb\xbf\x00 \t\r\n").lower()
    return any(
        marker in normalized[:8192]
        for marker in (b"<!doctype html", b"<html", b"<head", b"<meta")
    )


def extract_document(path, max_chars=30000):
    path = Path(path)
    extension = path.suffix.lower()
    with path.open("rb") as stream:
        prefix = stream.read(8192)
    if extension in {".pdf", ".docx", ".xlsx", ".xls"} and looks_like_html(prefix):
        return extract_html(path, max_chars)
    if extension == ".docx" and prefix.startswith(
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    ):
        return extract_doc(path, max_chars)

    extractors = {
        ".pdf": extract_pdf,
        ".xlsx": extract_xlsx,
        ".xls": extract_xls,
        ".docx": extract_docx,
        ".doc": extract_doc,
        ".csv": extract_csv,
    }
    extractor = extractors.get(extension)
    if not extractor:
        raise ExtractionError(f"unsupported extension: {extension}")
    return extractor(path, max_chars)
